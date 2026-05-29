import os
import re
import tempfile

import streamlit as st
from PyPDF2 import PdfReader
from docx import Document


def clean_xml_text(text: str) -> str:
    """
    PDF에서 추출된 텍스트 중 DOCX/XML에 저장할 수 없는 제어 문자를 제거합니다.
    예: NULL byte, 일부 control character 등
    """
    if text is None:
        return ""

    return re.sub(
        r"[\x00-\x08\x0B\x0C\x0E-\x1F]",
        "",
        text
    )


def pdf_to_docx_simple(pdf_path: str, docx_path: str):
    """
    PDF에서 텍스트만 추출하여 DOCX 파일로 저장합니다.
    원본 PDF의 이미지, 표, 레이아웃은 보존하지 않습니다.
    """
    reader = PdfReader(pdf_path, strict=False)
    doc = Document()

    for i, page in enumerate(reader.pages):
        text = page.extract_text()

        if text:
            text = clean_xml_text(text)

            for line in text.splitlines():
                line = clean_xml_text(line)

                if line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph()

        if i < len(reader.pages) - 1:
            doc.add_page_break()

    doc.save(docx_path)


def main():
    st.set_page_config(
        page_title="PDF → DOCX 변환기",
        page_icon="📄",
        layout="centered"
    )

    st.title("📄 PDF를 DOCX로 변환하기")
    st.write(
        "PDF 파일에서 텍스트만 추출하여 DOCX 파일로 변환합니다. "
        "이미지, 표, 세부 레이아웃은 보존되지 않습니다."
    )

    uploaded_file = st.file_uploader(
        "PDF 파일을 업로드하세요",
        type=["pdf"]
    )

    if uploaded_file is not None:
        st.info(f"선택된 파일: {uploaded_file.name}")

    if st.button("변환 시작"):
        if uploaded_file is None:
            st.warning("먼저 PDF 파일을 업로드해주세요.")
            return

        pdf_path = None
        docx_path = None

        with st.spinner("PDF를 DOCX로 변환하는 중입니다..."):
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                    tmp_pdf.write(uploaded_file.getvalue())
                    pdf_path = tmp_pdf.name

                base_name = os.path.splitext(os.path.basename(uploaded_file.name))[0]

                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                    docx_path = tmp_docx.name

                pdf_to_docx_simple(pdf_path, docx_path)

                with open(docx_path, "rb") as f:
                    docx_data = f.read()

                st.success("변환이 완료되었습니다!")

                st.download_button(
                    label="DOCX 파일 다운로드",
                    data=docx_data,
                    file_name=f"{base_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"변환 중 오류가 발생했습니다: {e}")

            finally:
                if pdf_path and os.path.exists(pdf_path):
                    os.remove(pdf_path)

                if docx_path and os.path.exists(docx_path):
                    os.remove(docx_path)


if __name__ == "__main__":
    main()
