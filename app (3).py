import streamlit as st
import io
from pypdf import PdfReader
from docx import Document

# 페이지 설정
st.set_page_config(
    page_title="PDF to DOCX 변환기",
    page_icon="📄",
    layout="centered"
)

st.title("📄 PDF to DOCX 변환기")
st.write("PDF 파일을 업로드하시면 텍스트를 추출하여 Word(.docx) 파일로 변환해 드립니다.")

# 파일 업로드 컴포넌트
uploaded_file = st.file_uploader("변환할 PDF 파일을 선택하세요.", type=["pdf"])

if uploaded_file is not None:
    try:
        with st.spinner("PDF 파일에서 텍스트를 분석하고 있습니다..."):
            # PDF 읽기
            pdf_reader = PdfReader(uploaded_file)
            doc = Document()
            
            full_text = []
            # 페이지별 텍스트 추출 및 Word 문서에 추가
            for i, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
                    full_text.append(text)
                else:
                    doc.add_paragraph(f"[페이지 {i+1}: 추출된 텍스트가 없거나 이미지 형식의 페이지입니다.]")
            
            # 추출된 텍스트가 아예 없는 경우 처리
            if not "".join(full_text).strip():
                st.warning("⚠️ PDF에서 추출된 텍스트가 없습니다. 스캔된 이미지형 PDF일 가능성이 있습니다.")
            
            # 메모리 버퍼에 Word 파일 저장
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            # 다운로드 버튼 활성화
            st.success("🎉 변환이 완료되었습니다!")
            
            # 원본 파일명에서 확장자 제외하고 .docx 붙이기
            original_name = uploaded_file.name
            output_name = original_name.rsplit(".", 1)[0] + ".docx"
            
            st.download_button(
                label="📥 변환된 Word 파일 다운로드",
                data=docx_buffer,
                file_name=output_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
    except Exception as e:
        st.error(f"❌ 변환 중 오류가 발생했습니다: {e}")
        st.info("파일이 손상되었거나 지원하지 않는 PDF 포맷일 수 있습니다.")
